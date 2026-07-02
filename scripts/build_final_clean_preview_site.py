#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from __future__ import annotations

import os
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from pypdf import PdfReader, PdfWriter

ROOT = Path(__file__).resolve().parents[1]
SRC_A = ROOT / "files/middle-school/grade-8/uncategorized/עבודת-קיץ-ח-אפשר-גם-אחרת-א.docx"
SRC_B = ROOT / "files/middle-school/grade-8/uncategorized/עבודת-קיץ-ח-אפשר-גם-אחרת-ב.docx"

OUT = ROOT / "previews/final-clean-h8"
HTML = ROOT / "previews/final-clean-h8.html"
PDF = OUT / "final-clean-h8.pdf"

TITLE_LINES = [
    "עבודת קיץ במתמטיקה",
    "הקבצה מדעית | שכבת ח׳",
    "תיכון סדיקול פסגת זאב",
]

FORBIDDEN = [
    "חלק א",
    "חלק ב",
    "שם התלמיד",
    "שם תלמיד",
    "שם:",
    "כיתה:",
    "כתה:",
    "תאריך:",
    "תרגילי חזרה לחופשת הקיץ",
    "על פי הספר אפשר גם אחרת",
]

REMOVE_RE = re.compile(
    r"עבודת\s+קיץ|"
    r"להקבצה\s+מדעית|"
    r"תיכון\s+סדיקול|"
    r"שכבת\s+ח|"
    r"חלק\s*[אב](?:׳|’)?|"
    r"שם\s*(?:התלמיד(?:/ה)?|תלמיד)?\s*:|"
    r"כיתה\s*:|"
    r"כתה\s*:|"
    r"תאריך\s*:|"
    r"תרגילי\s+חזרה\s+לחופשת\s+הקיץ|"
    r"על\s+פי\s+הספר\s+אפשר\s+גם\s+אחרת"
)

def norm(value: str) -> str:
    return " ".join((value or "").replace("\u00a0", " ").split())

def remove_xml(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)

def clear_part(part) -> None:
    for paragraph in list(part.paragraphs):
        remove_xml(paragraph._element)
    for table in list(part.tables):
        remove_xml(table._element)

def set_rtl_center(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if bidi is None:
        from docx.oxml import OxmlElement
        bidi = OxmlElement("w:bidi")
        ppr.append(bidi)
    bidi.set(qn("w:val"), "1")

def set_font(run, size: int, bold: bool) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold

def clean_doc(src: Path, dst: Path, add_title: bool) -> None:
    if not src.exists():
        raise FileNotFoundError(src)

    doc = Document(src)

    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.4)
        section.left_margin = Cm(1.45)
        section.right_margin = Cm(1.45)
        section.header_distance = Cm(0.1)
        section.footer_distance = Cm(0.1)
        section.different_first_page_header_footer = False

        for attr in [
            "header",
            "footer",
            "first_page_header",
            "first_page_footer",
            "even_page_header",
            "even_page_footer",
        ]:
            part = getattr(section, attr, None)
            if part is not None:
                if hasattr(part, "is_linked_to_previous"):
                    part.is_linked_to_previous = False
                clear_part(part)

    for index, paragraph in enumerate(list(doc.paragraphs)):
        text = norm(paragraph.text)
        if text and REMOVE_RE.search(text) and (index < 28 or len(text) <= 190):
            remove_xml(paragraph._element)

    for paragraph in list(doc.paragraphs[:18]):
        if norm(paragraph.text):
            break
        remove_xml(paragraph._element)

    if add_title:
        first = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        specs = [
            (TITLE_LINES[0], 22, True),
            (TITLE_LINES[1], 14, True),
            (TITLE_LINES[2], 13, True),
            ("", 8, False),
        ]
        for text, size, bold in specs:
            title_paragraph = first.insert_paragraph_before()
            set_rtl_center(title_paragraph)
            run = title_paragraph.add_run(text)
            set_font(run, size, bold)

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dst)

def convert(docx_path: Path) -> Path:
    soffice = os.environ.get("SOFFICE_EXE", "soffice")
    env = os.environ.copy()
    env["HOME"] = str(ROOT / ".lo-home-final-clean")
    Path(env["HOME"]).mkdir(exist_ok=True)

    subprocess.run(
        [
            soffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            str(OUT),
            str(docx_path),
        ],
        cwd=str(ROOT),
        env=env,
        check=True,
    )

    pdf_path = OUT / f"{docx_path.stem}.pdf"
    if not pdf_path.exists() or pdf_path.stat().st_size == 0:
        raise RuntimeError("PDF conversion failed")
    return pdf_path

def merge(files: list[Path]) -> None:
    writer = PdfWriter()
    for file_path in files:
        reader = PdfReader(str(file_path))
        for page in reader.pages:
            writer.add_page(page)

    writer.add_metadata({
        "/Title": "עבודת קיץ במתמטיקה",
        "/Creator": "LibreOffice + python-docx + pypdf",
    })

    with PDF.open("wb") as handle:
        writer.write(handle)

def pdf_text() -> str:
    reader = PdfReader(str(PDF))
    return "\n".join(page.extract_text() or "" for page in reader.pages)

def verify() -> int:
    reader = PdfReader(str(PDF))
    pages = len(reader.pages)

    if pages < 1:
        raise RuntimeError("Empty PDF")

    for page in reader.pages:
        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        if not (590 <= width <= 600 and 837 <= height <= 847):
            raise RuntimeError(f"Not A4: {width} x {height}")

    body = pdf_text()
    title_count = body.count(TITLE_LINES[0])

    if title_count != 1:
        raise RuntimeError(f"Title count must be 1, got {title_count}")

    hits = [item for item in FORBIDDEN if item in body]
    if hits:
        raise RuntimeError("Forbidden text still exists: " + ", ".join(hits))

    return pages

def write_html(pages: int) -> None:
    HTML.parent.mkdir(parents=True, exist_ok=True)
    html = """<!doctype html>
<html lang="he" dir="rtl">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>עבודת קיץ במתמטיקה</title>
<style>
body{margin:0;background:#eef3f8;font-family:Arial,sans-serif;color:#0f172a}
header{background:white;padding:16px 20px;border-bottom:1px solid #cbd5e1;position:sticky;top:0}
main{max-width:1100px;margin:auto;padding:18px}
h1{margin:0;font-size:24px}
.box{background:white;border:1px solid #cbd5e1;border-radius:16px;padding:12px;box-shadow:0 16px 45px #0002}
iframe{width:100%;height:92vh;border:0;border-radius:10px}
</style>
</head>
<body>
<header>
<h1>עבודת קיץ במתמטיקה</h1>
<div>PDF סופי: כותרת אחת בלבד בראש החוברת, ללא חלק א/ב וללא שדות תלמיד</div>
</header>
<main>
<div class="box"><iframe src="final-clean-h8/final-clean-h8.pdf"></iframe></div>
</main>
</body>
</html>"""
    HTML.write_text(html, encoding="utf-8")

def main() -> None:
    if OUT.exists():
        shutil.rmtree(OUT)

    OUT.mkdir(parents=True, exist_ok=True)

    part_a = OUT / "part-a-title-fixed.docx"
    part_b = OUT / "part-b-clean.docx"

    clean_doc(SRC_A, part_a, True)
    clean_doc(SRC_B, part_b, False)

    merge([convert(part_a), convert(part_b)])

    pages = verify()
    write_html(pages)

    print("FINAL_CLEAN_PDF=", PDF)
    print("PAGES=", pages)
    print("CHECK=OK: exactly one title, no part labels, no student fields, A4")

if __name__ == "__main__":
    main()
