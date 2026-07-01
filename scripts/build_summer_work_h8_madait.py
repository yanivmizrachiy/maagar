#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the real grade-8 scientific-track summer work from the original DOCX files.

The script edits the actual source Word files in the repository, adds a professional
Hebrew header/footer/title area, converts to PDF, merges the two parts, and creates
an external preview HTML page. It does not generate dummy body content.
"""
from __future__ import annotations

import os
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfWriter, PdfReader

ROOT = Path(__file__).resolve().parents[1]
SRC_A = ROOT / "files/middle-school/grade-8/uncategorized/עבודת-קיץ-ח-אפשר-גם-אחרת-א.docx"
SRC_B = ROOT / "files/middle-school/grade-8/uncategorized/עבודת-קיץ-ח-אפשר-גם-אחרת-ב.docx"
OUT = ROOT / "outputs/summer-work-h8-madait"
PREVIEWS = OUT / "preview-pages"
HTML_PATH = ROOT / "previews/avoda-kayitz-h8-madait-real.html"
TITLE_MAIN = "עבודת קיץ"
TITLE_SUB = "להקבצה מדעית | שכבת ח׳ | תיכון סדיקול פסגת זאב"
FOOTER = "עבודת קיץ • הקבצה מדעית • שכבת ח׳ • תיכון סדיקול פסגת זאב"


def set_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_run_font(run, size=None, bold=False, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    run.font.bold = bool(bold)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_cell_border(cell, color="CBD5E1", size="8"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def create_logo_placeholder(path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (700, 700), "white")
    d = ImageDraw.Draw(img)
    blue = (11, 79, 138)
    light = (226, 239, 250)
    d.ellipse((40, 40, 660, 660), fill=light, outline=blue, width=16)
    d.rectangle((185, 255, 515, 455), fill="white", outline=blue, width=10)
    d.polygon([(350, 145), (170, 260), (530, 260)], fill=(255, 255, 255), outline=blue)
    d.line((250, 455, 250, 325), fill=blue, width=10)
    d.line((350, 455, 350, 325), fill=blue, width=10)
    d.line((450, 455, 450, 325), fill=blue, width=10)
    try:
        font_big = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 58)
        font_small = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 42)
    except Exception:
        font_big = font_small = None
    d.text((350, 525), "סמל", anchor="mm", fill=blue, font=font_big)
    d.text((350, 585), "בית הספר", anchor="mm", fill=blue, font=font_small)
    img.save(path)


def clear_header_footer(container):
    for p in list(container.paragraphs):
        p._element.getparent().remove(p._element)
    for t in list(container.tables):
        t._element.getparent().remove(t._element)


def add_header(section, logo_path: Path, part_label: str):
    header = section.header
    header.is_linked_to_previous = False
    clear_header_footer(header)
    table = header.add_table(rows=1, cols=3, width=Cm(17.0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(3.0), Cm(10.8), Cm(3.0)]
    for i, width in enumerate(widths):
        table.columns[i].width = width
        cell = table.cell(0, i)
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "F8FAFC")
        add_cell_border(cell, "D8E3F0", "6")
    logo_cell = table.cell(0, 0)
    p_logo = logo_cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p_logo.add_run().add_picture(str(logo_path), width=Cm(1.65))
    except Exception:
        r = p_logo.add_run("סמל בית הספר")
        set_run_font(r, size=9, bold=True, color=(11, 79, 138))
    text_cell = table.cell(0, 1)
    p = text_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(TITLE_MAIN + "\n")
    set_run_font(r1, size=16, bold=True, color=(11, 79, 138))
    r2 = p.add_run(TITLE_SUB + "\n")
    set_run_font(r2, size=10.5, bold=True, color=(30, 41, 59))
    r3 = p.add_run(part_label)
    set_run_font(r3, size=9.5, bold=True, color=(71, 85, 105))
    empty = table.cell(0, 2).paragraphs[0]
    empty.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = empty.add_run("שם: ________\nכיתה: ______")
    set_run_font(r, size=8.5, bold=False, color=(71, 85, 105))


def add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    clear_header_footer(footer)
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(FOOTER)
    set_run_font(r, size=9, bold=True, color=(71, 85, 105))
    # add simple page field after separator
    r2 = p.add_run("  |  עמוד ")
    set_run_font(r2, size=9, bold=False, color=(100, 116, 139))
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=9, color=(100, 116, 139))


def insert_title_block(doc: Document, logo_path: Path, part_label: str):
    # Insert a real title block at the top without deleting any exercise content.
    first = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    title_p = first.insert_paragraph_before()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(TITLE_MAIN)
    set_run_font(r, size=22, bold=True, color=(11, 79, 138))
    sub_p = first.insert_paragraph_before()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub_p.add_run(TITLE_SUB)
    set_run_font(r, size=13, bold=True, color=(30, 41, 59))
    part_p = first.insert_paragraph_before()
    part_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = part_p.add_run(part_label)
    set_run_font(r, size=12, bold=True, color=(71, 85, 105))
    meta_p = first.insert_paragraph_before()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta_p.add_run("שם התלמיד/ה: ____________     כיתה: ______     תאריך: ______")
    set_run_font(r, size=11, bold=False, color=(71, 85, 105))
    # Divider paragraph
    div = first.insert_paragraph_before()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = div.add_run("────────────────────────────────────────")
    set_run_font(r, size=10, bold=False, color=(148, 163, 184))


def process_docx(src: Path, dst: Path, part_label: str, logo_path: Path):
    if not src.exists():
        raise FileNotFoundError(src)
    doc = Document(src)
    for section in doc.sections:
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1.7)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.header_distance = Cm(0.6)
        section.footer_distance = Cm(0.55)
        add_header(section, logo_path, part_label)
        add_footer(section)
    insert_title_block(doc, logo_path, part_label)
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dst)


def convert_to_pdf(docx_path: Path, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    env = os.environ.copy()
    env["HOME"] = str(ROOT / ".lo-home")
    (ROOT / ".lo-home").mkdir(exist_ok=True)
    cmd = [
        "libreoffice", "--headless", "--nologo", "--nofirststartwizard",
        "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)
    ]
    subprocess.run(cmd, check=True, cwd=str(ROOT), env=env)
    pdf = out_dir / (docx_path.stem + ".pdf")
    if not pdf.exists() or pdf.stat().st_size == 0:
        raise RuntimeError(f"PDF conversion failed for {docx_path}")
    return pdf


def merge_pdfs(paths, out_pdf: Path):
    writer = PdfWriter()
    for path in paths:
        reader = PdfReader(str(path))
        for page in reader.pages:
            writer.add_page(page)
    with out_pdf.open("wb") as f:
        writer.write(f)


def render_preview(pdf: Path):
    PREVIEWS.mkdir(parents=True, exist_ok=True)
    subprocess.run([
        "pdftoppm", "-png", "-f", "1", "-l", "4", "-r", "144", str(pdf), str(PREVIEWS / "page")
    ], check=True)


def write_html():
    HTML_PATH.parent.mkdir(parents=True, exist_ok=True)
    pdf_url = "https://github.com/yanivmizrachiy/maagar/raw/main/outputs/summer-work-h8-madait/avoda-kayitz-h8-madait-full.pdf"
    html = f"""<!doctype html>
<html lang=\"he\" dir=\"rtl\"><head><meta charset=\"utf-8\"><meta name=\"viewport\" content=\"width=device-width,initial-scale=1\">
<title>עבודת קיץ - הקבצה מדעית - שכבת ח׳</title>
<style>body{{margin:0;background:#f4f7fb;color:#0f172a;font-family:Arial,system-ui,sans-serif}}header{{background:#fff;border-bottom:1px solid #dbe4f0;padding:18px 22px;position:sticky;top:0;z-index:5}}main{{max-width:1120px;margin:auto;padding:22px}}h1{{margin:0;color:#0b4f8a;font-size:34px}}.sub{{margin:8px 0 0;color:#475569;font-size:18px}}.card{{background:white;border:1px solid #dbe4f0;border-radius:20px;box-shadow:0 18px 45px #0f172a20;padding:16px;margin-top:18px}}.actions{{display:flex;gap:10px;flex-wrap:wrap;margin-top:14px}}.btn{{display:inline-block;background:#0b4f8a;color:white;text-decoration:none;border-radius:999px;padding:11px 16px;font-weight:700}}.btn.green{{background:#0f766e}}.pdf{{width:100%;height:86vh;border:1px solid #cbd5e1;border-radius:16px;background:#fff}}</style>
</head><body><header><h1>עבודת קיץ</h1><div class=\"sub\">להקבצה מדעית · שכבת ח׳ · תיכון סדיקול פסגת זאב</div><div class=\"actions\"><a class=\"btn\" href=\"{pdf_url}\" target=\"_blank\" rel=\"noopener\">פתיחת PDF מלא</a><a class=\"btn green\" href=\"{pdf_url}\" download>הורדת PDF</a></div></header><main><section class=\"card\"><iframe class=\"pdf\" src=\"{pdf_url}\"></iframe></section></main></body></html>"""
    HTML_PATH.write_text(html, encoding="utf-8")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    logo = OUT / "school-logo-placeholder.png"
    create_logo_placeholder(logo)
    dst_a = OUT / "part-a-real-edited.docx"
    dst_b = OUT / "part-b-real-edited.docx"
    process_docx(SRC_A, dst_a, "חלק א׳", logo)
    process_docx(SRC_B, dst_b, "חלק ב׳", logo)
    pdf_a = convert_to_pdf(dst_a, OUT)
    pdf_b = convert_to_pdf(dst_b, OUT)
    full_pdf = OUT / "avoda-kayitz-h8-madait-full.pdf"
    merge_pdfs([pdf_a, pdf_b], full_pdf)
    render_preview(full_pdf)
    write_html()
    print("REAL_SOURCE_A=", SRC_A)
    print("REAL_SOURCE_B=", SRC_B)
    print("FULL_PDF=", full_pdf)


if __name__ == "__main__":
    main()
