#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from __future__ import annotations

import os
import re
import shutil
import subprocess
from pathlib import Path

import fitz
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pypdf import PdfReader, PdfWriter

ROOT = Path(__file__).resolve().parents[1]

SRC_A = ROOT / "files/middle-school/grade-8/uncategorized/עבודת-קיץ-ח-אפשר-גם-אחרת-א.docx"
SRC_B = ROOT / "files/middle-school/grade-8/uncategorized/עבודת-קיץ-ח-אפשר-גם-אחרת-ב.docx"

OUT = ROOT / "previews/final-clean-h8"
PDF = OUT / "final-clean-h8.pdf"
HTML = ROOT / "previews/final-clean-h8.html"

TITLE_LINES = [
    "עבודת קיץ במתמטיקה",
    "הקבצה מדעית | שכבת ח׳",
    "תיכון פסגת זאב ע״ש טדי קולק",
]

SUBMISSION_LINES = [
    "כל מי שלמד השנה בכיתה ח׳1 ונכח בקורס צריך להגיש חצי מהעבודה - סעיף כן וסעיף לא.",
    "מי שלא למד בכיתה ח׳1 השנה או לא נכח בקורס צריך לעשות את העבודה כולה.",
    "את העבודה יש להגיש בפתרון מלא במחברת קיץ.",
]

TOPIC_MAP = {
    "הפונקציה הקווית": "פונקציה קווית",
    "פונקציה קווית": "פונקציה קווית",
    "אי שוויונות": "אי שוויונות",
    "יחס": "יחס",
    "אחוזים": "אחוזים",
    "משוואות ממעלה ראשונה ושאלות מילוליות": "משוואות ממעלה ראשונה ושאלות מילוליות",
    "טכניקה אלגברית": "טכניקה אלגברית",
    "תיכון במשולש": "תיכון במשולש",
    "משולש שווה שוקיים": "משולש שווה שוקיים",
    "סטטיסטיקה": "סטטיסטיקה",
    "הסתברות": "הסתברות",
    "מערכת של שתי משוואות בשני נעלמים": "מערכת של שתי משוואות בשני נעלמים",
    "שאלות מילוליות בנושאים שונים": "שאלות מילוליות בנושאים שונים",
    "יחס ישר ויחס הפוך": "יחס ישר ויחס הפוך",
    "שורש ריבועי": "שורש ריבועי",
    "דמיון משולשים ודמיון מצולעים": "דמיון משולשים ודמיון מצולעים",
    "משפט פיתגורס": "משפט פיתגורס",
    "הגליל": "הגליל",
    "עבודה נעימה": "עבודה נעימה",
}

FORBIDDEN = [
    "חלק א",
    "חלק ב",
    "שם התלמיד",
    "שם תלמיד",
    "שם:",
    "כיתה:",
    "כתה:",
    "תאריך:",
    "הערה",
    "תרשים דירה",
    "מה אורך מרפסת",
    "צילמו תעודה",
    "קנה המידה של ההגדלה",
    "קנה המידה של ההקטנה",
    "גובה החרמון",
    "מודל של ההר",
    "מה גובה המודל",
]

REMOVE_TOP_RE = re.compile(
    r"עבודת\s+קיץ|"
    r"להקבצה\s+מדעית|"
    r"תיכון\s+סדיקול|"
    r"תיכון\s+פסגת\s+זאב|"
    r"שכבת\s+ח|"
    r"חלק\s*[אב](?:׳|’)?|"
    r"שם\s*(?:התלמיד(?:/ה)?|תלמיד)?\s*:|"
    r"כיתה\s*:|"
    r"כתה\s*:|"
    r"תאריך\s*:|"
    r"תרגילי\s+חזרה\s+לחופשת\s+הקיץ|"
    r"על\s+פי\s+הספר\s+אפשר\s+גם\s+אחרת"
)

def norm(value: str) -> str:
    return " ".join((value or "").replace("\u00a0", " ").split())

def remove_xml(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)

def clear_part(part) -> None:
    for paragraph in list(part.paragraphs):
        remove_xml(paragraph._element)
    for table in list(part.tables):
        remove_xml(table._element)

def set_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT) -> None:
    paragraph.alignment = align
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        ppr.append(bidi)
    bidi.set(qn("w:val"), "1")

def set_font(run, size: float, bold: bool = False, color: tuple[int, int, int] | None = None, name: str = "Arial") -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    rpr.rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def clear_runs_and_set(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)

def is_topic(text: str) -> str | None:
    t = norm(text)
    return TOPIC_MAP.get(t)

def style_topic(paragraph, text: str) -> None:
    clear_runs_and_set(paragraph, text)
    set_rtl(paragraph, WD_ALIGN_PARAGRAPH.RIGHT)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        set_font(run, 15, bold=True, color=(18, 83, 45))

def style_normal(paragraph) -> None:
    set_rtl(paragraph, WD_ALIGN_PARAGRAPH.RIGHT)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        if run.text:
            set_font(run, 11, bold=bool(run.bold), color=(0, 0, 0))

def remove_manual_page_breaks(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for br in list(run._element.iter(qn("w:br"))):
                if br.get(qn("w:type")) == "page":
                    parent = br.getparent()
                    if parent is not None:
                        parent.remove(br)

def remove_extra_empty_paragraphs(doc: Document) -> None:
    streak = 0
    for paragraph in list(doc.paragraphs):
        if norm(paragraph.text):
            streak = 0
            continue
        streak += 1
        if streak > 1:
            remove_xml(paragraph._element)

def remove_scale_section(doc: Document) -> None:
    start = None
    end = None
    paragraphs = list(doc.paragraphs)

    for i, p in enumerate(paragraphs):
        if norm(p.text) == "קנה מידה":
            start = i
            break

    if start is None:
        return

    for j in range(start + 1, len(paragraphs)):
        if "משוואות ממעלה ראשונה ושאלות מילוליות" in norm(paragraphs[j].text):
            end = j
            break

    if end is None:
        raise RuntimeError("נמצא קנה מידה אבל לא נמצא סוף החלק הבא: משוואות ממעלה ראשונה ושאלות מילוליות")

    for p in paragraphs[start:end]:
        remove_xml(p._element)

def clean_headers_footers(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.15)
        section.bottom_margin = Cm(1.15)
        section.left_margin = Cm(1.25)
        section.right_margin = Cm(1.25)
        section.header_distance = Cm(0.1)
        section.footer_distance = Cm(0.1)
        section.start_type = WD_SECTION_START.CONTINUOUS
        section.different_first_page_header_footer = False

        for attr in [
            "header",
            "footer",
            "first_page_header",
            "first_page_footer",
            "even_page_header",
            "even_page_footer",
        ]:
            part = getattr(section, attr, None)
            if part is not None:
                if hasattr(part, "is_linked_to_previous"):
                    part.is_linked_to_previous = False
                clear_part(part)

def remove_old_top_titles(doc: Document) -> None:
    for index, paragraph in enumerate(list(doc.paragraphs)):
        text = norm(paragraph.text)
        if text and REMOVE_TOP_RE.search(text) and (index < 30 or len(text) <= 190):
            remove_xml(paragraph._element)

def insert_opening_block(doc: Document) -> None:
    first = None
    for p in doc.paragraphs:
        if norm(p.text):
            first = p
            break

    if first is None:
        first = doc.add_paragraph()

    items = [
        ("", 2, False, (0, 0, 0), 0, 2),
        (SUBMISSION_LINES[2], 11.4, True, (160, 20, 35), 0, 2),
        (SUBMISSION_LINES[1], 11.4, True, (160, 20, 35), 0, 1),
        (SUBMISSION_LINES[0], 11.4, True, (160, 20, 35), 0, 5),
        ("", 2, False, (0, 0, 0), 0, 3),
        (TITLE_LINES[2], 14.2, True, (12, 45, 95), 0, 2),
        (TITLE_LINES[1], 13.3, True, (22, 65, 125), 0, 1),
        (TITLE_LINES[0], 19.0, True, (12, 45, 95), 0, 2),
    ]

    for text, size, bold, color, before, after in items:
        p = first.insert_paragraph_before()
        set_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        set_font(r, size, bold=bold, color=color)

def normalize_document(doc: Document, add_title: bool, remove_scale: bool) -> None:
    clean_headers_footers(doc)
    remove_manual_page_breaks(doc)
    remove_old_top_titles(doc)

    if remove_scale:
        remove_scale_section(doc)

    remove_extra_empty_paragraphs(doc)

    if add_title:
        insert_opening_block(doc)

    for paragraph in doc.paragraphs:
        text = norm(paragraph.text)
        if not text:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            continue

        topic = is_topic(text)
        if topic:
            style_topic(paragraph, topic)
        else:
            style_normal(paragraph)

def convert_to_pdf(docx_path: Path) -> Path:
    soffice = os.environ.get("SOFFICE_EXE", "soffice")
    env = os.environ.copy()
    env["HOME"] = str(ROOT / ".lo-home-premium-reflow")
    Path(env["HOME"]).mkdir(exist_ok=True)

    subprocess.run(
        [
            soffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            str(OUT),
            str(docx_path),
        ],
        cwd=str(ROOT),
        env=env,
        check=True,
    )

    pdf_path = OUT / f"{docx_path.stem}.pdf"
    if not pdf_path.exists() or pdf_path.stat().st_size == 0:
        raise RuntimeError("PDF conversion failed")
    return pdf_path

def merge_pdfs(items: list[Path]) -> None:
    writer = PdfWriter()

    for item in items:
        reader = PdfReader(str(item))
        for page in reader.pages:
            writer.add_page(page)

    writer.add_metadata({
        "/Title": "עבודת קיץ במתמטיקה",
        "/Creator": "premium DOCX reflow + LibreOffice + pypdf + PyMuPDF",
    })

    with PDF.open("wb") as handle:
        writer.write(handle)

def pdf_text() -> str:
    reader = PdfReader(str(PDF))
    return "\n".join(page.extract_text() or "" for page in reader.pages)

def validate_pdf() -> None:
    reader = PdfReader(str(PDF))
    pages = len(reader.pages)

    if pages < 1:
        raise RuntimeError("PDF ריק")

    for page in reader.pages:
        w = float(page.mediabox.width)
        h = float(page.mediabox.height)
        if not (590 <= w <= 600 and 837 <= h <= 847):
            raise RuntimeError(f"עמוד לא A4: {w} x {h}")

    text = pdf_text()
    title_count = text.count("עבודת קיץ במתמטיקה")

    if title_count != 1:
        raise RuntimeError(f"כותרת ראשית חייבת להופיע פעם אחת בלבד. בפועל: {title_count}")

    hits = [x for x in FORBIDDEN if x in text]
    if hits:
        raise RuntimeError("נשארו טקסטים אסורים: " + ", ".join(hits))

    required = [
        "פונקציה קווית",
        "אי שוויונות",
        "יחס",
        "משוואות ממעלה ראשונה ושאלות מילוליות",
        "סטטיסטיקה",
        "משפט פיתגורס",
    ]
    missing = [x for x in required if x not in text]
    if missing:
        raise RuntimeError("חסרות כותרות נושא: " + ", ".join(missing))

    fitz_doc = fitz.open(str(PDF))
    too_empty = []
    for index, page in enumerate(fitz_doc, start=1):
        blocks = page.get_text("blocks")
        real = [b for b in blocks if len((b[4] or "").strip()) > 10]
        if not real:
            too_empty.append(index)

    if too_empty:
        raise RuntimeError("עמודים ריקים/חלשים מדי: " + ", ".join(map(str, too_empty)))

def write_html() -> None:
    html = """<!doctype html>
<html lang="he" dir="rtl">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>עבודת קיץ במתמטיקה</title>
<style>
body{margin:0;background:#eef3f8;font-family:Arial,sans-serif;color:#0f172a}
header{background:white;padding:16px 20px;border-bottom:1px solid #cbd5e1;position:sticky;top:0}
main{max-width:1100px;margin:auto;padding:18px}
h1{margin:0;font-size:24px;color:#0b2d5f}
.box{background:white;border:1px solid #cbd5e1;border-radius:16px;padding:12px;box-shadow:0 16px 45px #0002}
iframe{width:100%;height:92vh;border:0;border-radius:10px}
</style>
</head>
<body>
<header>
<h1>עבודת קיץ במתמטיקה</h1>
<div>גרסה נקייה: עימוד מחדש מהמקור, ללא קנה מידה, ללא כותרות כפולות, ללא קווים מיותרים.</div>
</header>
<main>
<div class="box"><iframe src="final-clean-h8/final-clean-h8.pdf"></iframe></div>
</main>
</body>
</html>"""
    HTML.parent.mkdir(parents=True, exist_ok=True)
    HTML.write_text(html, encoding="utf-8")

def render_previews() -> None:
    previews = OUT / "qa-previews"
    if previews.exists():
        shutil.rmtree(previews)
    previews.mkdir(parents=True, exist_ok=True)

    doc = fitz.open(str(PDF))
    wanted = [1, 3, 6, 14, len(doc)]

    for page_number in wanted:
        if 1 <= page_number <= len(doc):
            page = doc[page_number - 1]
            pix = page.get_pixmap(matrix=fitz.Matrix(1.7, 1.7), alpha=False)
            pix.save(previews / f"page-{page_number:02d}.png")

def main() -> None:
    if OUT.exists():
        shutil.rmtree(OUT)
    OUT.mkdir(parents=True, exist_ok=True)

    part_a_docx = OUT / "part-a-premium-reflow.docx"
    part_b_docx = OUT / "part-b-premium-reflow.docx"

    doc_a = Document(SRC_A)
    normalize_document(doc_a, add_title=True, remove_scale=True)
    doc_a.save(part_a_docx)

    doc_b = Document(SRC_B)
    normalize_document(doc_b, add_title=False, remove_scale=False)
    doc_b.save(part_b_docx)

    part_a_pdf = convert_to_pdf(part_a_docx)
    part_b_pdf = convert_to_pdf(part_b_docx)

    merge_pdfs([part_a_pdf, part_b_pdf])
    validate_pdf()
    write_html()
    render_previews()

    print("FINAL_PDF=", PDF)
    print("PAGES=", len(PdfReader(str(PDF)).pages))
    print("CHECK=OK: rebuilt from DOCX, no scale section, one title, no forbidden text, A4, no empty pages")

if __name__ == "__main__":
    main()


